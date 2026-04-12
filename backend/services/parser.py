import os
import uuid
import hashlib
from io import BytesIO
from typing import BinaryIO

from backend.schemas.resume import ParsedResume, ResumeSection


ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def allowed_file(filename: str) -> bool:
    _, ext = os.path.splitext(filename.lower())
    return ext in ALLOWED_EXTENSIONS


def get_file_type(filename: str) -> str:
    _, ext = os.path.splitext(filename.lower())
    return ext[1:]  # Remove the dot


def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(BytesIO(file_content))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)
    except ImportError:
        return ""


def extract_text_from_docx(file_content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(file_content))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except ImportError:
        return ""


def extract_text(filename: str, file_content: bytes) -> str:
    file_type = get_file_type(filename)
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_content)
    elif file_type == "docx":
        return extract_text_from_docx(file_content)
    else:
        return file_content.decode("utf-8", errors="ignore")


def generate_file_hash(file_content: bytes) -> str:
    return hashlib.sha256(file_content).hexdigest()


def parse_resume_text(text: str) -> ParsedResume:
    sections = parse_sections(text)
    
    return ParsedResume(
        contact=extract_contact(text),
        summary=sections.get("summary", [None])[0] if sections.get("summary") else None,
        experience=sections.get("experience", []),
        education=sections.get("education", []),
        skills=sections.get("skills", []),
        certifications=sections.get("certifications", []),
        projects=sections.get("projects", []),
        links=extract_links(text),
    )


def parse_sections(text: str) -> dict[str, list]:
    lines = text.split("\n")
    sections: dict[str, list[ResumeSection | str]] = {}
    current_section = "other"
    current_items: list[str] = []
    
    section_keywords = [
        ("summary", ["summary", "objective", "profile"]),
        ("experience", ["experience", "work history", "employment", "professional experience"]),
        ("education", ["education", "academic", "degree", "university", "college"]),
        ("skills", ["skills", "technical skills", "competencies", "technologies"]),
        ("certifications", ["certification", "certifications", "license", "licenses"]),
        ("projects", ["projects", "personal projects", "portfolio"]),
    ]
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        is_section_header = False
        for section_name, keywords in section_keywords:
            if any(line_stripped.lower() == keyword.lower() for keyword in keywords) or \
               any(line_stripped.lower().startswith(keyword.lower() + ":") for keyword in keywords):
                if current_items:
                    if current_section == "skills":
                        sections[current_section] = current_items
                    else:
                        sections.setdefault(current_section, []).extend(current_items)
                current_section = section_name
                current_items = []
                is_section_header = True
                break
        
        if not is_section_header:
            current_items.append(line_stripped)
    
    if current_items:
        if current_section == "skills":
            sections[current_section] = current_items
        elif current_section != "other":
            sections.setdefault(current_section, []).extend(current_items)
        else:
            sections.setdefault("other", []).extend(current_items)
    
    parsed_sections = {}
    for key, value in sections.items():
        if key == "skills":
            parsed_sections[key] = value
        elif key == "summary":
            parsed_sections[key] = value
        else:
            parsed_sections[key] = [
                ResumeSection(description=item) if isinstance(item, str) else item
                for item in value
            ]
    
    return parsed_sections


def extract_contact(text: str) -> dict | None:
    lines = text.split("\n")[:10]
    contact = {}
    
    for line in lines:
        line = line.strip()
        if "@" in line and "." in line:
            email_parts = line.split()
            for part in email_parts:
                if "@" in part and "." in part.split("@")[-1]:
                    contact["email"] = part
                    break
        
        phone_chars = [c for c in line if c.isdigit() or c in "()-+ "]
        if len(phone_chars) >= 10:
            contact["phone"] = "".join(phone_chars).strip()
    
    return contact if contact else None


def extract_links(text: str) -> list[str]:
    import re
    url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
    urls = re.findall(url_pattern, text)
    return list(set(urls))
